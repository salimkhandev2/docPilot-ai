from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_assignment_docx():
    # Create a new Document
    doc = Document()
    
    # Set default font to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Assignment 2 - Software Testing Techniques', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Add spacing
    doc.add_paragraph()
    
    # ========== SECTION 1: Decision Table Testing ==========
    section1_title = doc.add_heading('1. Decision Table Testing', 1)
    section1_title_run = section1_title.runs[0]
    section1_title_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark Green
    section1_title_run.font.bold = True
    
    # System description
    doc.add_paragraph('System: Online Course Discount Eligibility System', style='Heading 3')
    
    # Conditions section
    conditions_heading = doc.add_heading('Conditions:', 3)
    conditions_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)  # Firebrick
    
    # Create bullet list for conditions
    conditions = [
        ("C1: User Type", "(Student, Educator, General Public)"),
        ("C2: Course Category", "(Technical, Non-Technical)"),
        ("C3: Enrollment Time", "(Early Bird, Regular, Last Minute)"),
        ("C4: Previous Purchases", "(First-time Buyer, Returning Customer)")
    ]
    
    for condition, values in conditions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(condition).bold = True
        p.add_run(f" {values}")
    
    # Actions section
    doc.add_paragraph()
    actions_heading = doc.add_heading('Actions/Outcomes:', 3)
    actions_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    actions = [
        "A1: Apply 50% discount",
        "A2: Apply 30% discount", 
        "A3: Apply 20% discount",
        "A4: Apply 10% discount",
        "A5: No discount"
    ]
    
    for action in actions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(action).bold = True
    
    # Decision Table section
    doc.add_paragraph()
    table_heading = doc.add_heading('Decision Table:', 3)
    table_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    # Create the decision table
    table = doc.add_table(rows=37, cols=6)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Rule No.', 'C1: User Type', 'C2: Course Type', 'C3: Enrollment Time', 
               'C4: Previous Purchase', 'Actions']
    
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        # Set background color for header
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '4F81BD')  # Blue background
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Decision table data
    decision_data = [
        [1, 'Student', 'Technical', 'Early Bird', 'First-time', 'A1'],
        [2, 'Student', 'Technical', 'Early Bird', 'Returning', 'A2'],
        [3, 'Student', 'Technical', 'Regular', 'First-time', 'A2'],
        [4, 'Student', 'Technical', 'Regular', 'Returning', 'A3'],
        [5, 'Student', 'Technical', 'Last Minute', 'First-time', 'A4'],
        [6, 'Student', 'Technical', 'Last Minute', 'Returning', 'A4'],
        [7, 'Student', 'Non-Technical', 'Early Bird', 'First-time', 'A2'],
        [8, 'Student', 'Non-Technical', 'Early Bird', 'Returning', 'A3'],
        [9, 'Student', 'Non-Technical', 'Regular', 'First-time', 'A3'],
        [10, 'Student', 'Non-Technical', 'Regular', 'Returning', 'A4'],
        [11, 'Student', 'Non-Technical', 'Last Minute', 'First-time', 'A4'],
        [12, 'Student', 'Non-Technical', 'Last Minute', 'Returning', 'A5'],
        [13, 'Educator', 'Technical', 'Early Bird', 'First-time', 'A1'],
        [14, 'Educator', 'Technical', 'Early Bird', 'Returning', 'A1'],
        [15, 'Educator', 'Technical', 'Regular', 'First-time', 'A2'],
        [16, 'Educator', 'Technical', 'Regular', 'Returning', 'A2'],
        [17, 'Educator', 'Technical', 'Last Minute', 'First-time', 'A3'],
        [18, 'Educator', 'Technical', 'Last Minute', 'Returning', 'A3'],
        [19, 'Educator', 'Non-Technical', 'Early Bird', 'First-time', 'A2'],
        [20, 'Educator', 'Non-Technical', 'Early Bird', 'Returning', 'A2'],
        [21, 'Educator', 'Non-Technical', 'Regular', 'First-time', 'A3'],
        [22, 'Educator', 'Non-Technical', 'Regular', 'Returning', 'A3'],
        [23, 'Educator', 'Non-Technical', 'Last Minute', 'First-time', 'A4'],
        [24, 'Educator', 'Non-Technical', 'Last Minute', 'Returning', 'A4'],
        [25, 'General', 'Technical', 'Early Bird', 'First-time', 'A3'],
        [26, 'General', 'Technical', 'Early Bird', 'Returning', 'A3'],
        [27, 'General', 'Technical', 'Regular', 'First-time', 'A4'],
        [28, 'General', 'Technical', 'Regular', 'Returning', 'A4'],
        [29, 'General', 'Technical', 'Last Minute', 'First-time', 'A5'],
        [30, 'General', 'Technical', 'Last Minute', 'Returning', 'A5'],
        [31, 'General', 'Non-Technical', 'Early Bird', 'First-time', 'A4'],
        [32, 'General', 'Non-Technical', 'Early Bird', 'Returning', 'A4'],
        [33, 'General', 'Non-Technical', 'Regular', 'First-time', 'A5'],
        [34, 'General', 'Non-Technical', 'Regular', 'Returning', 'A5'],
        [35, 'General', 'Non-Technical', 'Last Minute', 'First-time', 'A5'],
        [36, 'General', 'Non-Technical', 'Last Minute', 'Returning', 'A5']
    ]
    
    # Fill table with data
    for i, row_data in enumerate(decision_data, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_data)
    
    # Test Cases section
    doc.add_paragraph()
    test_cases_heading = doc.add_heading('Test Cases Derived:', 3)
    test_cases_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    test_cases = [
        {
            'name': 'Test Case 1 (Rule 1):',
            'input': 'User Type = Student, Course Type = Technical, Enrollment Time = Early Bird, Previous Purchase = First-time',
            'expected': '50% discount',
            'purpose': 'Test maximum discount scenario for students'
        },
        {
            'name': 'Test Case 2 (Rule 12):',
            'input': 'User Type = Student, Course Type = Non-Technical, Enrollment Time = Last Minute, Previous Purchase = Returning',
            'expected': 'No discount',
            'purpose': 'Test no discount scenario'
        },
        {
            'name': 'Test Case 3 (Rule 14):',
            'input': 'User Type = Educator, Course Type = Technical, Enrollment Time = Early Bird, Previous Purchase = Returning',
            'expected': '50% discount',
            'purpose': 'Test educator benefits'
        },
        {
            'name': 'Test Case 4 (Rule 29):',
            'input': 'User Type = General, Course Type = Technical, Enrollment Time = Last Minute, Previous Purchase = First-time',
            'expected': 'No discount',
            'purpose': 'Test general public with unfavorable conditions'
        },
        {
            'name': 'Test Case 5 (Rule 20):',
            'input': 'User Type = Educator, Course Type = Non-Technical, Enrollment Time = Early Bird, Previous Purchase = Returning',
            'expected': '30% discount',
            'purpose': 'Test non-technical course discount for educators'
        }
    ]
    
    for i, tc in enumerate(test_cases, 1):
        p = doc.add_paragraph()
        p.add_run(tc['name']).bold = True
        p.add_run('\nInput: ' + tc['input'])
        p.add_run('\nExpected: ').bold = True
        p.add_run(tc['expected'])
        p.add_run('\nPurpose: ').bold = True
        p.add_run(tc['purpose'])
        if i < len(test_cases):
            doc.add_paragraph()  # Add spacing between test cases
    
    # Coverage section
    doc.add_paragraph()
    coverage_heading = doc.add_heading('How Decision Table Testing Ensures Complete Coverage:', 3)
    coverage_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    coverage_points = [
        "Exhaustive Combination Testing: The decision table systematically covers all possible combinations of conditions (3 × 2 × 3 × 2 = 36 combinations), ensuring no business rule is missed.",
        "Business Rule Validation: Each column represents a specific business rule, making it easy to verify that each rule is implemented correctly.",
        "Elimination of Redundancy: By structuring all conditions and outcomes, we can identify and eliminate redundant test cases.",
        "Consistency Checking: The table format helps identify contradictory rules where similar conditions lead to different outcomes.",
        "Traceability: Each test case is directly traceable to specific business rules, making it easy to demonstrate test coverage to stakeholders."
    ]
    
    for i, point in enumerate(coverage_points, 1):
        p = doc.add_paragraph(style='List Number')
        parts = point.split(':')
        p.add_run(parts[0] + ':').bold = True
        if len(parts) > 1:
            p.add_run(parts[1])
    
    # Add page break
    doc.add_page_break()
    
    # ========== SECTION 2: State Transition Testing ==========
    section2_title = doc.add_heading('2. State Transition Testing', 1)
    section2_title_run = section2_title.runs[0]
    section2_title_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark Green
    section2_title_run.font.bold = True
    
    # Application description
    doc.add_paragraph('Application: ATM Banking System', style='Heading 3')
    
    # States section
    states_heading = doc.add_heading('States:', 3)
    states_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    states = [
        "Idle State (S1): ATM ready for card insertion",
        "Card Inserted (S2): Card validated, waiting for PIN",
        "PIN Entered (S3): PIN verified, showing main menu",
        "Balance Inquiry (S4): Displaying account balance",
        "Cash Withdrawal (S5): Processing withdrawal request",
        "Transaction Complete (S6): Transaction finalized, card returned",
        "Invalid PIN (S7): PIN entry failed",
        "Card Retained (S8): Card kept by machine due to security",
        "Out of Service (S9): ATM temporarily unavailable"
    ]
    
    for state in states:
        p = doc.add_paragraph(style='List Bullet')
        parts = state.split(':')
        p.add_run(parts[0]).bold = True
        if len(parts) > 1:
            p.add_run(':' + parts[1])
    
    # Transitions table
    doc.add_paragraph()
    transitions_heading = doc.add_heading('Transitions and Events:', 3)
    transitions_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    # Create transitions table
    trans_table = doc.add_table(rows=22, cols=3)
    trans_table.style = 'Table Grid'
    
    # Transitions header
    trans_headers = ['Current State', 'Event', 'Next State']
    for i, header in enumerate(trans_headers):
        cell = trans_table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C0504D')  # Red background
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Transitions data
    transitions = [
        ['S1: Idle', 'E1: Insert Valid Card', 'S2'],
        ['S1: Idle', 'E2: Insert Invalid Card', 'S9'],
        ['S2: Card Inserted', 'E3: Enter Correct PIN', 'S3'],
        ['S2: Card Inserted', 'E4: Enter Wrong PIN (1st attempt)', 'S7'],
        ['S2: Card Inserted', 'E5: Cancel Transaction', 'S6'],
        ['S3: PIN Entered', 'E6: Select Balance Inquiry', 'S4'],
        ['S3: PIN Entered', 'E7: Select Cash Withdrawal', 'S5'],
        ['S3: PIN Entered', 'E8: Select Cancel', 'S6'],
        ['S4: Balance Inquiry', 'E9: Request Another Service', 'S3'],
        ['S4: Balance Inquiry', 'E10: Finish Transaction', 'S6'],
        ['S5: Cash Withdrawal', 'E11: Sufficient Funds', 'S6'],
        ['S5: Cash Withdrawal', 'E12: Insufficient Funds', 'S3'],
        ['S5: Cash Withdrawal', 'E13: Cancel Withdrawal', 'S3'],
        ['S6: Transaction Complete', 'E14: Take Card', 'S1'],
        ['S6: Transaction Complete', 'E15: Timeout (card not taken)', 'S8'],
        ['S7: Invalid PIN', 'E16: Enter Correct PIN (2nd attempt)', 'S3'],
        ['S7: Invalid PIN', 'E17: Enter Wrong PIN (2nd attempt)', 'S7'],
        ['S7: Invalid PIN', 'E18: Enter Wrong PIN (3rd attempt)', 'S8'],
        ['S7: Invalid PIN', 'E19: Cancel', 'S6'],
        ['S8: Card Retained', 'E20: System Reset', 'S1'],
        ['S9: Out of Service', 'E21: System Recovery', 'S1']
    ]
    
    # Fill transitions table
    for i, transition in enumerate(transitions, 1):
        for j, data in enumerate(transition):
            cell = trans_table.cell(i, j)
            cell.text = data
    
    # State Transition Diagram
    doc.add_paragraph()
    diagram_heading = doc.add_heading('State Transition Diagram Representation:', 3)
    diagram_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    diagram_lines = [
        'S1 → E1 → S2 → E3 → S3 → E6 → S4 → E10 → S6 → E14 → S1',
        'S1 → E2 → S9 → E21 → S1',
        'S2 → E4 → S7 → E16 → S3',
        'S2 → E5 → S6',
        'S3 → E7 → S5 → E11 → S6',
        'S3 → E8 → S6',
        'S5 → E12 → S3',
        'S5 → E13 → S3',
        'S6 → E15 → S8 → E20 → S1',
        'S7 → E17 → S7',
        'S7 → E18 → S8',
        'S7 → E19 → S6'
    ]
    
    for line in diagram_lines:
        p = doc.add_paragraph(line)
        p.style = 'Normal'
    
    # Test Scenarios section
    doc.add_paragraph()
    scenarios_heading = doc.add_heading('Test Scenarios:', 3)
    scenarios_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    # Valid Scenarios
    valid_scenarios_heading = doc.add_heading('Valid Transition Scenarios:', 4)
    valid_scenarios_heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
    
    valid_scenarios = [
        {
            'name': '1. Happy Path - Successful Withdrawal:',
            'path': 'S1 → E1 → S2 → E3 → S3 → E7 → S5 → E11 → S6 → E14 → S1',
            'test': 'Insert valid card → Enter correct PIN → Select withdrawal → Sufficient funds → Complete → Take card'
        },
        {
            'name': '2. Balance Inquiry:',
            'path': 'S1 → E1 → S2 → E3 → S3 → E6 → S4 → E9 → S3 → E8 → S6 → E14 → S1',
            'test': 'Check balance → Return to menu → Cancel transaction'
        },
        {
            'name': '3. PIN Recovery:',
            'path': 'S1 → E1 → S2 → E4 → S7 → E16 → S3 → E8 → S6 → E14 → S1',
            'test': 'Wrong PIN first attempt → Correct PIN second attempt → Cancel'
        }
    ]
    
    for scenario in valid_scenarios:
        p = doc.add_paragraph()
        p.add_run(scenario['name']).bold = True
        p.add_run('\nPath: ' + scenario['path'])
        p.add_run('\nTest: ' + scenario['test'])
        doc.add_paragraph()
    
    # Invalid Scenarios
    invalid_scenarios_heading = doc.add_heading('Invalid Transition Scenarios:', 4)
    invalid_scenarios_heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
    
    invalid_scenarios = [
        {
            'name': '1. Card Security - Multiple PIN Failures:',
            'path': 'S1 → E1 → S2 → E4 → S7 → E17 → S7 → E18 → S8',
            'test': 'Three wrong PIN attempts → Card retained'
        },
        {
            'name': '2. Timeout Scenario:',
            'path': 'S1 → E1 → S2 → E3 → S3 → E8 → S6 → [Wait timeout] → S8',
            'test': 'Complete transaction but don\'t take card → Card retained'
        },
        {
            'name': '3. Invalid Card Insertion:',
            'path': 'S1 → E2 → S9',
            'test': 'Insert damaged/invalid card → System out of service'
        },
        {
            'name': '4. Insufficient Funds:',
            'path': 'S1 → E1 → S2 → E3 → S3 → E7 → S5 → E12 → S3 → E8 → S6 → E14 → S1',
            'test': 'Attempt withdrawal with insufficient funds → Return to menu → Cancel'
        }
    ]
    
    for scenario in invalid_scenarios:
        p = doc.add_paragraph()
        p.add_run(scenario['name']).bold = True
        p.add_run('\nPath: ' + scenario['path'])
        p.add_run('\nTest: ' + scenario['test'])
        if scenario != invalid_scenarios[-1]:
            doc.add_paragraph()
    
    # Benefits section
    doc.add_paragraph()
    benefits_heading = doc.add_heading('How State Transition Testing Helps Uncover Errors:', 3)
    benefits_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)
    
    benefits_points = [
        "Invalid State Detection: Tests reveal if the system enters undefined or illegal states (e.g., allowing withdrawal without PIN verification).",
        "Transition Errors: Identifies incorrect state changes (e.g., moving from S7 to S1 instead of S8 after three PIN failures).",
        "Event Handling Issues: Uncovers problems with event processing (e.g., system not responding to timeout events).",
        "Concurrency Problems: Helps identify race conditions when multiple events occur simultaneously.",
        "Memory Issues: Detects problems with state persistence (e.g., system 'forgetting' it has a card inserted).",
        "Boundary Condition Testing: Specifically tests edge cases like maximum failed attempts, timeouts, and recovery scenarios.",
        "Recovery Testing: Validates that the system can recover correctly from error states to normal operation.",
        "Sequence Dependency: Identifies issues where the order of operations affects system behavior."
    ]
    
    for i, point in enumerate(benefits_points, 1):
        p = doc.add_paragraph(style='List Number')
        parts = point.split(':')
        p.add_run(parts[0] + ':').bold = True
        if len(parts) > 1:
            p.add_run(parts[1])
    
    # Conclusion paragraph
    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.add_run('Conclusion: ').bold = True
    conclusion.add_run('This approach is particularly effective for systems like ATMs where security, reliability, and correct sequencing of operations are critical. It ensures that all possible user interactions are tested, including both normal flows and exceptional cases that might otherwise be overlooked.')
    
    # Save the document
    filename = 'Assignment2_Software_Testing_Techniques.docx'
    doc.save(filename)
    print(f"Document saved as: {filename}")
    print(f"File location: {os.path.abspath(filename)}")
    
    return filename

# Main execution
if __name__ == "__main__":
    # Install required packages if needed
    required_packages = ['python-docx']
    
    print("Creating Assignment 2 Document...")
    try:
        filename = create_assignment_docx()
        print("\n" + "="*60)
        print("SUCCESS: Document created successfully!")
        print("="*60)
        print(f"\nThe document '{filename}' has been created with:")
        print("✓ Arial font throughout")
        print("✓ Colored headings and sections")
        print("✓ Formatted tables with colored headers")
        print("✓ Complete Assignment 2 content")
        print("✓ Professional formatting and structure")
        print("\nYou can now open the document in Microsoft Word.")
    except ImportError as e:
        print(f"Error: {e}")
        print("\nPlease install required packages using:")
        print("pip install python-docx")
    except Exception as e:
        print(f"An error occurred: {e}")